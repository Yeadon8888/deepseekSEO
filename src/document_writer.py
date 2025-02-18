from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from utils.config import Config
import os
import logging
from datetime import datetime

class DocumentWriter:
    def __init__(self):
        self.config = Config()
        self.logger = logging.getLogger(__name__)
        self.output_dir = self.config.get('output_dir', 'output')
        
    def create_document(self, content, image_paths):
        """
        创建Word文档并添加内容和图片
        :param content: 文章内容
        :param image_paths: 图片路径列表
        :return: 生成的文档路径
        """
        try:
            doc = Document()
            
            # 设置文档样式
            self._set_document_style(doc)
            
            # 处理文章内容和图片
            self._add_content_with_images(doc, content, image_paths)
            
            # 保存文档
            if not os.path.exists(self.output_dir):
                os.makedirs(self.output_dir)
                
            # 使用时间戳生成文件名
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = os.path.join(self.output_dir, f'article_{timestamp}.docx')
            doc.save(filename)
            
            self.logger.info(f"文档已保存到: {filename}")
            return filename
            
        except Exception as e:
            error_msg = f"生成文档失败: {str(e)}"
            self.logger.error(error_msg)
            raise Exception(error_msg)
    
    def _set_document_style(self, doc):
        """
        设置文档的基本样式
        """
        # 设置默认字体
        style = doc.styles['Normal']
        style.font.name = '微软雅黑'
        style.font.size = Pt(12)
        
        # 设置标题样式
        for i in range(1, 4):  # h1 to h3
            style = doc.styles[f'Heading {i}']
            style.font.name = '微软雅黑'
            style.font.size = Pt(20 - i * 2)  # h1: 18pt, h2: 16pt, h3: 14pt
            style.font.bold = True
            style.font.color.rgb = RGBColor(0, 0, 0)
            
        # 设置图片说明样式
        if 'Caption' not in doc.styles:
            style = doc.styles.add_style('Caption', 1)
        else:
            style = doc.styles['Caption']
        style.font.name = '微软雅黑'
        style.font.size = Pt(10)
        style.font.italic = True
        style.font.color.rgb = RGBColor(89, 89, 89)
    
    def _add_content_with_images(self, doc, content, image_paths):
        """
        添加文章内容和图片，实现图文混排
        """
        # 移除META标签
        content = self._remove_meta_tags(content)
        
        # 按段落分割内容
        paragraphs = content.split('\n\n')
        image_index = 0
        
        # 添加标题
        if paragraphs and paragraphs[0].strip().startswith('#'):
            title_text = ' '.join(paragraphs[0].split()[1:])
            heading = doc.add_heading(title_text, 0)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraphs = paragraphs[1:]
        
        # 处理正文内容
        for i, para in enumerate(paragraphs):
            if not para.strip():
                continue
                
            if para.strip().startswith('#'):  # 标题
                level = len(para.split()[0]) - 1
                text = ' '.join(para.split()[1:])
                doc.add_heading(text, level=min(level, 3))
            else:  # 普通段落
                doc.add_paragraph(para.strip())
            
            # 每隔3个段落插入一张图片（如果还有图片的话）
            if i > 0 and i % 3 == 0 and image_index < len(image_paths):
                self._insert_image_with_caption(
                    doc,
                    image_paths[image_index],
                    f"图{image_index + 1}",
                    image_index
                )
                image_index += 1
        
        # 插入剩余的图片
        while image_index < len(image_paths):
            self._insert_image_with_caption(
                doc,
                image_paths[image_index],
                f"图{image_index + 1}",
                image_index
            )
            image_index += 1
    
    def _insert_image_with_caption(self, doc, image_path, caption_text, index):
        """
        插入图片和说明文字
        """
        if not os.path.exists(image_path):
            self.logger.warning(f"图片文件不存在: {image_path}")
            return
            
        try:
            # 添加分隔段落
            doc.add_paragraph()
            
            # 添加图片
            picture = doc.add_picture(image_path, width=Inches(6.0))
            
            # 居中对齐图片
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 添加图片说明
            caption = doc.add_paragraph(style='Caption')
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption.add_run(caption_text).italic = True
            
            # 添加空行
            doc.add_paragraph()
            
        except Exception as e:
            self.logger.error(f"插入图片失败: {str(e)}")
    
    def _remove_meta_tags(self, content):
        """
        移除SEO META标签
        """
        lines = content.split('\n')
        start_idx = -1
        end_idx = -1
        
        for i, line in enumerate(lines):
            if '<!--SEO Meta Tags-->' in line:
                start_idx = i
            elif '<!--End SEO Meta Tags-->' in line:
                end_idx = i
                break
        
        if start_idx != -1 and end_idx != -1:
            return '\n'.join(lines[:start_idx] + lines[end_idx + 1:])
        return content